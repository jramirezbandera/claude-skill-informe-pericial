"""Inyecta el presupuesto de Presto/Arquímedes (.RTF) en el informe v1.

Pre-requisitos:
- `<encargo>/99_salida/informe_pericial_v1.docx` generado por redactar_v1.py
  (con marcadores `[[PRESUPUESTO_DEFICIENCIA: <slug>]]`, `[[COSTE]]`,
  `[[HOJA_RESUMEN_PRESUPUESTO ...]]`)
- `<encargo>/99_salida/presupuesto/*.rtf` exportado de Presto/Arquímedes.
  Si hay varios RTF, se usa el más reciente por mtime.

Mapeo deficiencia ↔ subcapítulo:
- Cada deficiencia en `caso.yaml > deficiencias` puede llevar `subcapitulo_presupuesto: "04.05"`.
- Si falta, se hace match fuzzy entre el slug y los títulos de subcapítulos del RTF.
- Si tampoco hay match, se inserta marcador `[[SUBCAPÍTULO NO ENCONTRADO]]` y se avisa.

Salida:
- `<encargo>/99_salida/informe_pericial_v2.docx`

Uso:
    python presupuestar.py --encargo "<ruta>"
"""
from __future__ import annotations

import argparse
import re
import sys
import unicodedata
from copy import deepcopy
from decimal import Decimal
from pathlib import Path

import yaml
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Cm, Pt
from striprtf.striprtf import rtf_to_text

try:
    sys.stdout.reconfigure(encoding="utf-8")
    sys.stderr.reconfigure(encoding="utf-8")
except Exception:
    pass

EXT_RTF = {".rtf"}


# ────────── Parser RTF de presupuesto Presto ──────────

RE_CAPITULO = re.compile(r"^\s*CAP[IÍ]TULO\s+(\d+)\s+(.+?)\s*$", re.IGNORECASE)
RE_SUBCAPITULO = re.compile(r"^\s*SUBCAP[IÍ]TULO\s+([\d.]+)\s+(.+?)\s*$", re.IGNORECASE)
RE_TOTAL_CAP = re.compile(r"^\s*TOTAL\s+CAP[IÍ]TULO\s+(\d+).*?([\d.]+,\d{2})\s*$", re.IGNORECASE)
RE_TOTAL_SUBCAP = re.compile(r"^\s*TOTAL\s+SUBCAP[IÍ]TULO\s+([\d.]+)", re.IGNORECASE)
RE_TOTAL_GENERAL = re.compile(r"^\s*TOTAL\s+([\d.]+,\d{2})\s*$", re.IGNORECASE)
# Cabecera de partida: "01.01.01\tm2\tDESCRIPCIÓN"
RE_PARTIDA_CAB = re.compile(r"^\s*([\d.]+\.[\d.]+\.\d+)\s+(\w+)\s+(.+?)\s*$")
# Línea de medición: cualquier línea con tabs y números (ej. "zonas ajardinadas\t1\t163,00\t0,90\t146,70")
RE_NUMEROS_TABBED = re.compile(r"^\s*[^\t]*\t.*\d", re.UNICODE)
# Importe en formato europeo: "8.111,25" o "1.705,68"
RE_IMPORTE = re.compile(r"(\d{1,3}(?:\.\d{3})*,\d{2})")


def _normalizar(s: str) -> str:
    """Slug-style: lower, sin acentos, sin no-alfanum."""
    s = unicodedata.normalize("NFKD", s).encode("ascii", "ignore").decode()
    return re.sub(r"[^a-z0-9]+", "", s.lower())


def _parse_importe(s: str) -> Decimal:
    """'8.111,25' -> Decimal('8111.25')."""
    if not s:
        return Decimal("0")
    s = s.strip().replace(".", "").replace(",", ".")
    try:
        return Decimal(s)
    except Exception:
        return Decimal("0")


def _fmt_eur(d: Decimal) -> str:
    """Decimal('8111.25') -> '8.111,25 €'."""
    s = f"{d:,.2f}"
    s = s.replace(",", "X").replace(".", ",").replace("X", ".")
    return f"{s} €"


def parsear_rtf_presto(rtf_path: Path) -> dict:
    """Parsea un RTF de Presto. Devuelve:
    {
      "capitulos": [
        {"codigo": "01", "titulo": "...", "subcapitulos": [...], "total": Decimal}
      ],
      "subcapitulos_por_codigo": {"01.01": {...}},
      "total_general": Decimal,
    }
    Cada subcapítulo: {codigo, titulo, partidas: [...], total: Decimal, lineas_brutas: [str]}
    """
    raw = rtf_path.read_bytes().decode("latin-1", errors="replace")
    txt = rtf_to_text(raw)
    lines = txt.splitlines()

    capitulos = []
    subcapitulos_idx = {}
    cap_actual = None
    sub_actual = None
    partida_actual = None
    total_general = Decimal("0")

    def flush_partida():
        nonlocal partida_actual
        if partida_actual and sub_actual is not None:
            sub_actual["partidas"].append(partida_actual)
        partida_actual = None

    def flush_sub():
        nonlocal sub_actual
        if sub_actual and cap_actual is not None:
            cap_actual["subcapitulos"].append(sub_actual)
            subcapitulos_idx[sub_actual["codigo"]] = sub_actual
        sub_actual = None

    def flush_cap():
        nonlocal cap_actual
        if cap_actual:
            capitulos.append(cap_actual)
        cap_actual = None

    # Para totales subcapítulo que vienen en línea siguiente
    pending_total_sub_codigo = None
    pending_total_sub_ref = None  # ref al dict del subcap (ya flusheado)

    def _cancel_pending():
        nonlocal pending_total_sub_codigo, pending_total_sub_ref
        pending_total_sub_codigo = None
        pending_total_sub_ref = None

    for line in lines:
        line_strip = line.strip()

        # Prioridad 1: TOTAL CAPÍTULO (anula pending)
        m_total_cap = RE_TOTAL_CAP.match(line)
        if m_total_cap:
            _cancel_pending()
            flush_partida()
            flush_sub()
            if cap_actual:
                cap_actual["total"] = _parse_importe(m_total_cap.group(2))
            flush_cap()
            continue

        # Prioridad 2: TOTAL GENERAL (línea aislada con sólo "TOTAL\t<importe>")
        m_total_gen = re.match(r"^\s*TOTAL\s+([\d.]+,\d{2})\s*$", line)
        if not m_total_gen:
            # también puede venir con tabs: "\tTOTAL\t\t44.357,14"
            m_total_gen = re.match(r"^\s*TOTAL[\s\t]+([\d.]+,\d{2})\s*$", line)
        if m_total_gen and not re.search(r"CAP[IÍ]TULO|SUBCAP", line, re.IGNORECASE):
            _cancel_pending()
            total_general = _parse_importe(m_total_gen.group(1))
            continue

        # Prioridad 3: TOTAL SUBCAPÍTULO
        m_total_sub = RE_TOTAL_SUBCAP.match(line)
        if m_total_sub:
            flush_partida()
            codigo = m_total_sub.group(1)
            imp_match = RE_IMPORTE.search(line)
            if sub_actual and sub_actual["codigo"] == codigo:
                if imp_match:
                    sub_actual["total"] = _parse_importe(imp_match.group(1))
                    _cancel_pending()
                else:
                    pending_total_sub_codigo = codigo
                    pending_total_sub_ref = sub_actual
            flush_sub()
            continue

        # Si hay un total subcap pendiente (importe en línea siguiente), buscarlo aquí
        if pending_total_sub_codigo:
            m_imp = RE_IMPORTE.search(line)
            # Pero sólo lo aceptamos si la línea es "continuación" — sin palabras clave
            if m_imp and not re.search(r"CAP[IÍ]TULO|SUBCAP", line, re.IGNORECASE):
                if pending_total_sub_ref is not None:
                    pending_total_sub_ref["total"] = _parse_importe(m_imp.group(1))
                _cancel_pending()
                continue
            # Si la línea es solo continuación de título sin importe, seguir buscando
            if line_strip and not re.search(r"CAP[IÍ]TULO|SUBCAP", line, re.IGNORECASE) and not RE_PARTIDA_CAB.match(line):
                continue
            # Llegamos a algo importante; abandonar y reprocesar
            _cancel_pending()

        m_cap = RE_CAPITULO.match(line)
        if m_cap:
            flush_partida()
            flush_sub()
            flush_cap()
            cap_actual = {
                "codigo": m_cap.group(1),
                "titulo": m_cap.group(2).strip(),
                "subcapitulos": [],
                "total": Decimal("0"),
            }
            continue

        m_sub = RE_SUBCAPITULO.match(line)
        if m_sub:
            flush_partida()
            flush_sub()
            sub_actual = {
                "codigo": m_sub.group(1),
                "titulo": m_sub.group(2).strip(),
                "partidas": [],
                "total": Decimal("0"),
                "lineas_brutas": [],
            }
            continue

        if sub_actual is None:
            continue

        sub_actual["lineas_brutas"].append(line)

        m_cab = RE_PARTIDA_CAB.match(line)
        if m_cab:
            flush_partida()
            partida_actual = {
                "codigo": m_cab.group(1),
                "ud": m_cab.group(2),
                "descripcion_breve": m_cab.group(3).strip(),
                "descripcion_larga": [],
                "mediciones": [],
                "cantidad": Decimal("0"),
                "precio": Decimal("0"),
                "importe": Decimal("0"),
            }
            continue

        if partida_actual is None:
            continue

        # Líneas con tabs: pueden ser mediciones o el total de partida
        if "\t" in line:
            partes = [p for p in line.split("\t")]
            partes_strip = [p.strip() for p in partes]
            # Línea de total de partida: empieza con tab vacío y tiene 3 números (cantidad, precio, importe)
            # ej. "\t\t231,75\t7,36\t1.705,68"
            num_partes = [p for p in partes_strip if p]
            if (not partes_strip[0]) and len(num_partes) >= 3 and all(re.match(r"^[\d.,]+$", p) for p in num_partes[-3:]):
                partida_actual["cantidad"] = _parse_importe(num_partes[-3])
                partida_actual["precio"] = _parse_importe(num_partes[-2])
                partida_actual["importe"] = _parse_importe(num_partes[-1])
                continue
            # Es una medición: descripción + uds + longitud + anchura + altura + parcial
            desc = partes_strip[0] if partes_strip else ""
            nums = [p for p in partes_strip[1:] if p]
            partida_actual["mediciones"].append({"desc": desc, "valores": nums})
            continue

        # Línea sin tabs: descripción larga
        if line.strip():
            partida_actual["descripcion_larga"].append(line.strip())

    flush_partida()
    flush_sub()
    flush_cap()

    return {
        "capitulos": capitulos,
        "subcapitulos_por_codigo": subcapitulos_idx,
        "total_general": total_general,
    }


# ────────── Mapeo deficiencia ↔ subcapítulo ──────────

def mapear_deficiencias(deficiencias: list[dict], presupuesto: dict, caso: dict) -> dict[str, dict]:
    """Devuelve {slug: subcapitulo_dict | None}.

    Prioridades:
    1. caso.yaml > deficiencias[i].subcapitulo_presupuesto (código tipo "04.05")
    2. caso.yaml > deficiencias[i].capitulo_presupuesto (código o título — backward compat)
    3. Match fuzzy: slug normalizado contra título de subcapítulo normalizado
    """
    out: dict[str, dict] = {}
    subcaps = presupuesto["subcapitulos_por_codigo"]
    # Diccionario auxiliar para fuzzy match
    norm_a_codigo = {_normalizar(s["titulo"]): s["codigo"] for s in subcaps.values()}

    # Mapeos manuales desde caso.yaml
    cy_defs = {d.get("slug"): d for d in caso.get("deficiencias", [])}

    for d in deficiencias:
        slug = d["slug"]
        cy = cy_defs.get(slug, {})
        codigo_explicito = cy.get("subcapitulo_presupuesto") or cy.get("capitulo_presupuesto")
        if codigo_explicito:
            # Puede ser código directo o texto que contenga el código
            m = re.search(r"\b(\d+\.\d+)\b", str(codigo_explicito))
            if m and m.group(1) in subcaps:
                out[slug] = subcaps[m.group(1)]
                continue
        # Fuzzy: slug normalizado debe estar contenido en el título normalizado
        slug_norm = _normalizar(slug)
        candidatos = []
        for sub in subcaps.values():
            tit_norm = _normalizar(sub["titulo"])
            if slug_norm in tit_norm or tit_norm in slug_norm:
                candidatos.append(sub)
        if len(candidatos) == 1:
            out[slug] = candidatos[0]
        elif len(candidatos) > 1:
            print(f"AVISO: '{slug}' tiene {len(candidatos)} matches en presupuesto, ninguno asignado: {[c['codigo'] for c in candidatos]}", file=sys.stderr)
            out[slug] = None
        else:
            out[slug] = None
    return out


# ────────── Construcción del bloque docx para un subcapítulo ──────────

def construir_bloque_subcapitulo(doc, subcap: dict, anchor_para) -> None:
    """Inserta antes de anchor_para el bloque del subcapítulo: tabla con partidas + total.

    El bloque sale como una tabla con 6 columnas:
    Código | Ud | Descripción | Cantidad | Precio (€) | Importe (€)
    """
    body = anchor_para._p.getparent()
    anchor_idx = list(body).index(anchor_para._p)

    def _insertar(elemento_xml):
        nonlocal anchor_idx
        body.insert(anchor_idx, elemento_xml)
        anchor_idx += 1

    # Encabezado del subcapítulo (párrafo en negrita)
    p = doc.add_paragraph()
    r = p.add_run(f"Subcapítulo {subcap['codigo']} — {subcap['titulo']}")
    r.bold = True
    r.font.size = Pt(11)
    _insertar(p._p)

    # Tabla con partidas
    tabla = doc.add_table(rows=1, cols=6)
    try:
        tabla.style = doc.styles["Table Grid"]
    except KeyError:
        pass
    hdr = tabla.rows[0].cells
    for i, h in enumerate(["Código", "Ud", "Descripción", "Cantidad", "Precio (€)", "Importe (€)"]):
        hdr[i].text = h
        for para in hdr[i].paragraphs:
            for run in para.runs:
                run.bold = True
                run.font.size = Pt(9)

    for partida in subcap["partidas"]:
        row = tabla.add_row().cells
        row[0].text = partida["codigo"]
        row[1].text = partida["ud"]
        # Descripción: breve + larga (resumida)
        desc = partida["descripcion_breve"]
        if partida["descripcion_larga"]:
            larga = " ".join(partida["descripcion_larga"])
            if len(larga) > 200:
                larga = larga[:197] + "..."
            desc = f"{desc}\n{larga}"
        row[2].text = desc
        row[3].text = f"{partida['cantidad']:.2f}".replace(".", ",")
        row[4].text = _fmt_eur(partida["precio"]).replace(" €", "")
        row[5].text = _fmt_eur(partida["importe"]).replace(" €", "")
        # Reducir font de toda la fila
        for cell in row:
            for para in cell.paragraphs:
                for run in para.runs:
                    run.font.size = Pt(8)
    _insertar(tabla._tbl)

    # Total subcapítulo (párrafo derecha)
    p_total = doc.add_paragraph()
    p_total.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r = p_total.add_run(f"TOTAL SUBCAPÍTULO {subcap['codigo']}: {_fmt_eur(subcap['total'])}")
    r.bold = True
    _insertar(p_total._p)

    # Espacio
    p_sep = doc.add_paragraph()
    _insertar(p_sep._p)


# ────────── Sustitución de marcadores ──────────

def encontrar_marcador(doc, marcador_regex: re.Pattern):
    """Devuelve generator de (paragraph, match)."""
    for para in doc.paragraphs:
        m = marcador_regex.search(para.text)
        if m:
            yield para, m


def eliminar_parrafo(para):
    """Elimina el párrafo del cuerpo del documento."""
    el = para._p
    el.getparent().remove(el)


def reemplazar_texto_en_celda(celda, antiguo: str, nuevo: str) -> bool:
    """True si reemplazó."""
    if antiguo in celda.text:
        # Recolectar todos los runs y rehacerlos
        for para in celda.paragraphs:
            for run in para.runs:
                if antiguo in run.text:
                    run.text = run.text.replace(antiguo, nuevo)
                    return True
            # El antiguo puede estar partido entre runs — caso raro pero posible
            full = "".join(r.text for r in para.runs)
            if antiguo in full:
                new_full = full.replace(antiguo, nuevo)
                # Vaciar runs y poner el nuevo en el primero
                for run in para.runs:
                    run.text = ""
                if para.runs:
                    para.runs[0].text = new_full
                return True
    return False


# ────────── Main ──────────

def main(argv: list[str]) -> int:
    p = argparse.ArgumentParser()
    p.add_argument("--encargo", required=True)
    args = p.parse_args(argv)

    encargo = Path(args.encargo).resolve()
    v1_path = encargo / "99_salida" / "informe_pericial_v1.docx"
    if not v1_path.exists():
        print(f"ERROR: no existe {v1_path}. Ejecuta /informe redactar antes.", file=sys.stderr)
        return 1

    pres_dir = encargo / "99_salida" / "presupuesto"
    if not pres_dir.exists():
        print(f"ERROR: no existe {pres_dir}", file=sys.stderr)
        return 1

    rtfs = sorted([p for p in pres_dir.iterdir()
                    if p.is_file() and p.suffix.lower() in EXT_RTF],
                   key=lambda p: p.stat().st_mtime, reverse=True)
    if not rtfs:
        print(f"ERROR: no hay .rtf en {pres_dir}", file=sys.stderr)
        return 1
    rtf = rtfs[0]
    print(f"Usando presupuesto: {rtf.name}")

    # caso.yaml
    caso = yaml.safe_load((encargo / "caso.yaml").read_text(encoding="utf-8")) or {}

    # deficiencias.md (lista canónica)
    def_path = encargo / "deficiencias.md"
    deficiencias = []
    for line in (def_path.read_text(encoding="utf-8") if def_path.exists() else "").splitlines():
        m = re.match(r"^\s*(\d+)\.\s+\*\*([a-z0-9_]+)\*\*\s*·\s*(.+?)\s*$", line)
        if m:
            deficiencias.append({"id": int(m.group(1)), "slug": m.group(2), "titulo": m.group(3).strip()})

    if not deficiencias:
        print(f"ERROR: deficiencias.md vacío.", file=sys.stderr)
        return 1

    # Parsear presupuesto
    presupuesto = parsear_rtf_presto(rtf)
    print(f"Capítulos detectados: {len(presupuesto['capitulos'])}")
    print(f"Subcapítulos detectados: {len(presupuesto['subcapitulos_por_codigo'])}")
    print(f"Total general detectado: {_fmt_eur(presupuesto['total_general'])}")

    # Mapear
    mapeo = mapear_deficiencias(deficiencias, presupuesto, caso)
    print("\nMapeo deficiencia → subcapítulo:")
    for slug, sub in mapeo.items():
        if sub:
            print(f"  {slug} → {sub['codigo']} ({sub['titulo'][:50]})  [{_fmt_eur(sub['total'])}]")
        else:
            print(f"  {slug} → NO MATCH")

    # Cargar v1
    doc = Document(str(v1_path))

    # Sustituir marcadores [[PRESUPUESTO_DEFICIENCIA: <slug>]]
    pat_pres = re.compile(r"\[\[PRESUPUESTO_DEFICIENCIA:\s*([a-z0-9_]+)\s*\]\]")
    parrafos_a_eliminar = []
    n_inyectados = 0
    n_no_encontrados = 0
    for para, m in encontrar_marcador(doc, pat_pres):
        slug = m.group(1)
        sub = mapeo.get(slug)
        if not sub:
            # Sustituir el marcador por aviso pero no eliminar
            for run in para.runs:
                if "[[PRESUPUESTO_DEFICIENCIA" in run.text:
                    run.text = re.sub(pat_pres, "[[SUBCAPÍTULO NO ENCONTRADO]]", run.text)
            n_no_encontrados += 1
            continue
        construir_bloque_subcapitulo(doc, sub, para)
        parrafos_a_eliminar.append(para)
        n_inyectados += 1
    for para in parrafos_a_eliminar:
        eliminar_parrafo(para)

    # Sustituir [[COSTE]] en cuadro resumen
    n_costes = 0
    for tabla in doc.tables:
        if not tabla.rows:
            continue
        # Detectar el cuadro resumen por la cabecera
        hdr_text = " | ".join(c.text.strip().lower() for c in tabla.rows[0].cells)
        if "deficiencia" in hdr_text and "coste" in hdr_text:
            for fila in tabla.rows[1:]:
                # La columna 1 (índice 1) es el título; usamos la columna 0 (id) o el título para mapear
                try:
                    id_str = fila.cells[0].text.strip()
                    titulo = fila.cells[1].text.strip()
                except IndexError:
                    continue
                # Encontrar el slug que casa
                slug = None
                for d in deficiencias:
                    if str(d["id"]) == id_str or d["titulo"] == titulo:
                        slug = d["slug"]
                        break
                if slug and mapeo.get(slug):
                    coste = _fmt_eur(mapeo[slug]["total"])
                    if reemplazar_texto_en_celda(fila.cells[-1], "[[COSTE]]", coste):
                        n_costes += 1

    # Sustituir [[HOJA_RESUMEN_PRESUPUESTO ...]] por una tabla con totales por capítulo
    pat_resumen = re.compile(r"\[\[HOJA_RESUMEN_PRESUPUESTO[^\]]*\]\]")
    parrafos_resumen = list(encontrar_marcador(doc, pat_resumen))
    for para, _ in parrafos_resumen:
        # Insertar una tabla con todos los capítulos y sus totales
        body = para._p.getparent()
        anchor_idx = list(body).index(para._p)

        tabla = doc.add_table(rows=1, cols=2)
        try:
            tabla.style = doc.styles["Table Grid"]
        except KeyError:
            pass
        hdr = tabla.rows[0].cells
        hdr[0].text = "Capítulo"
        hdr[1].text = "Importe (€)"
        for cell in hdr:
            for p_h in cell.paragraphs:
                for r_h in p_h.runs:
                    r_h.bold = True
        for cap in presupuesto["capitulos"]:
            row = tabla.add_row().cells
            row[0].text = f"{cap['codigo']}. {cap['titulo']}"
            row[1].text = _fmt_eur(cap["total"]).replace(" €", "")
            row[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
        # Fila total general
        row = tabla.add_row().cells
        row[0].text = "TOTAL GENERAL"
        row[1].text = _fmt_eur(presupuesto["total_general"]).replace(" €", "")
        for cell in row:
            for p_h in cell.paragraphs:
                for r_h in p_h.runs:
                    r_h.bold = True
        row[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT

        body.insert(anchor_idx, tabla._tbl)
        eliminar_parrafo(para)

    # Guardar v2
    out = encargo / "99_salida" / "informe_pericial_v2.docx"
    doc.save(str(out))

    print(f"\nOK — informe v2 generado: {out}")
    print(f"     Tamaño: {out.stat().st_size/1024:.0f} KB")
    print(f"     Subcapítulos inyectados: {n_inyectados}")
    if n_no_encontrados:
        print(f"     AVISO: {n_no_encontrados} deficiencias sin subcapítulo (marcador [[SUBCAPÍTULO NO ENCONTRADO]] en el docx)")
    print(f"     Costes en cuadro resumen: {n_costes}/{len(deficiencias)}")
    if parrafos_resumen:
        print(f"     Hoja resumen presupuesto: tabla con {len(presupuesto['capitulos'])} capítulos + total general")
    return 0


if __name__ == "__main__":
    sys.exit(main(sys.argv[1:]))
