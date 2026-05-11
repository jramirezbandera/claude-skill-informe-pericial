"""Recorre TODA la documentación del encargo (recursivamente) y extrae texto plano.

Procesa cualquier subcarpeta dentro de:
    00_encargo/
    01_documentacion/    (incluye subcarpetas como "Certificaciones y fras", "Correos relevantes"...)
    02_parte_contraria/
    03_comunicaciones/
    05_ensayos/
    06_referencias/

NO procesa: 04_inspeccion/notas_tablet/  (eso lo hace leer_notability.py)
NO procesa: 04_inspeccion/fotos*/          (las fotos se gestionan aparte)
NO procesa: 99_salida/                     (el output del informe)
NO procesa: _skill_workspace/              (workspace interno)

Para cada archivo encontrado:
- .pdf con texto extraíble → escribe `<archivo>.pdf.txt` al lado
- .pdf escaneado (sin texto) → renderiza preview de la primera página como `<archivo>.pdf.preview.png`
- .docx / .doc → escribe `<archivo>.docx.txt` al lado
- .txt / .md → ya legibles, sólo se listan
- imágenes → sólo se listan, no se extrae nada
- otros → se listan como "no extraíble"

Genera `_skill_workspace/inventario.md` con la tabla completa.

Uso:
    python extraer_textos.py --encargo "Z:\\03-INFORME\\Encargos\\..."
    python extraer_textos.py --encargo "..." --force      # re-extrae aunque ya exista .txt
"""
from __future__ import annotations

import argparse
import sys
from pathlib import Path

import fitz  # pymupdf

try:
    from docx import Document as DocxDocument
except ImportError:
    DocxDocument = None

try:
    sys.stdout.reconfigure(encoding="utf-8")
    sys.stderr.reconfigure(encoding="utf-8")
except Exception:
    pass


# Subcarpetas raíz que SÍ procesamos (recursivamente)
CARPETAS_PROCESAR = [
    "00_encargo",
    "01_documentacion",
    "02_parte_contraria",
    "03_comunicaciones",
    "05_ensayos",
    "06_referencias",
]

# Carpetas que NUNCA procesamos aunque estén dentro
EXCLUDED_DIRS = {"_skill_workspace", "_pages", "fotos", "fotos_renombradas", "notas_tablet", "99_salida", "04_inspeccion"}

MIN_TEXT_CHARS = 50

EXT_IMAGEN = {".jpg", ".jpeg", ".png", ".gif", ".webp", ".heic", ".heif", ".tif", ".tiff", ".bmp"}
EXT_VIDEO = {".mp4", ".mov", ".avi", ".mkv"}
EXT_AUDIO = {".mp3", ".wav", ".m4a", ".ogg"}


def es_archivo_generado(nombre: str) -> bool:
    n = nombre.lower()
    return n.endswith((".pdf.txt", ".pdf.preview.png", ".docx.txt", ".doc.txt"))


def extraer_pdf(pdf: Path, force: bool = False) -> tuple[str, int]:
    out_txt = pdf.with_suffix(pdf.suffix + ".txt")
    if out_txt.exists() and not force:
        return ("skip", out_txt.stat().st_size)
    try:
        doc = fitz.open(str(pdf))
    except Exception as e:
        return (f"err:{e}", 0)
    text_chars = 0
    bloques = []
    for i in range(doc.page_count):
        t = doc[i].get_text()
        text_chars += len(t)
        bloques.append(f"\n=== PÁGINA {i+1} ===\n{t}")
    if text_chars >= MIN_TEXT_CHARS:
        out_txt.write_text("".join(bloques), encoding="utf-8")
        return ("ok", text_chars)
    # Escaneado: render preview de la primera página
    preview = pdf.with_suffix(pdf.suffix + ".preview.png")
    if not preview.exists() or force:
        try:
            pix = doc[0].get_pixmap(dpi=150)
            pix.save(str(preview))
        except Exception:
            pass
    return ("scanned", text_chars)


def extraer_docx(docx: Path, force: bool = False) -> tuple[str, int]:
    out_txt = docx.with_suffix(docx.suffix + ".txt")
    if out_txt.exists() and not force:
        return ("skip", out_txt.stat().st_size)
    if DocxDocument is None:
        return ("err:python-docx no instalado", 0)
    try:
        d = DocxDocument(str(docx))
    except Exception as e:
        return (f"err:{e}", 0)
    bloques = []
    for para in d.paragraphs:
        if para.text.strip():
            bloques.append(para.text)
    for tbl in d.tables:
        for row in tbl.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    if para.text.strip():
                        bloques.append(para.text)
    txt = "\n".join(bloques)
    out_txt.write_text(txt, encoding="utf-8")
    return ("ok", len(txt))


def walk_files(carpeta: Path):
    """Recorre recursivamente, saltando carpetas excluidas."""
    for ruta in sorted(carpeta.rglob("*")):
        # Saltar si está dentro de una carpeta excluida (en cualquier nivel)
        if any(part in EXCLUDED_DIRS for part in ruta.parts):
            continue
        if ruta.is_file() and not ruta.name.startswith(".") and not es_archivo_generado(ruta.name):
            yield ruta


def categorizar(f: Path) -> str:
    ext = f.suffix.lower()
    if ext == ".pdf": return "PDF"
    if ext in (".docx", ".doc"): return "DOCX"
    if ext in (".txt", ".md"): return "TEXTO"
    if ext == ".yaml" or ext == ".yml": return "YAML"
    if ext in EXT_IMAGEN: return "IMAGEN"
    if ext in EXT_VIDEO: return "VIDEO"
    if ext in EXT_AUDIO: return "AUDIO"
    if ext in (".xls", ".xlsx", ".ods"): return "EXCEL"
    if ext in (".csv", ".tsv"): return "CSV"
    return ext.lstrip(".").upper() or "?"


def main(argv: list[str]) -> int:
    p = argparse.ArgumentParser()
    p.add_argument("--encargo", required=True)
    p.add_argument("--force", action="store_true")
    args = p.parse_args(argv)

    encargo = Path(args.encargo)
    if not encargo.exists():
        print(f"ERROR: no existe {encargo}", file=sys.stderr)
        return 1

    inventario = []  # (subcarpeta, ruta_relativa, kb, tipo, chars, estado)
    workspace = encargo / "_skill_workspace"
    workspace.mkdir(exist_ok=True)

    n_total = 0
    n_extraidos = 0
    for sub in CARPETAS_PROCESAR:
        carpeta = encargo / sub
        if not carpeta.exists():
            continue
        for f in walk_files(carpeta):
            n_total += 1
            rel = f.relative_to(encargo)
            kb = round(f.stat().st_size / 1024)
            tipo = categorizar(f)

            if tipo == "PDF":
                estado, chars = extraer_pdf(f, args.force)
                if estado in ("ok", "skip"): n_extraidos += 1
                tag = ("✓ texto" if estado == "ok" else
                       "⚠ escaneado (PDF imagen)" if estado == "scanned" else
                       "⏭ ya extraído" if estado == "skip" else
                       f"✗ {estado}")
                inventario.append((sub, str(rel), kb, "PDF", chars, tag))
            elif tipo == "DOCX":
                estado, chars = extraer_docx(f, args.force)
                if estado in ("ok", "skip"): n_extraidos += 1
                tag = ("✓ texto" if estado == "ok" else
                       "⏭ ya extraído" if estado == "skip" else
                       f"✗ {estado}")
                inventario.append((sub, str(rel), kb, "DOCX", chars, tag))
            elif tipo == "TEXTO":
                inventario.append((sub, str(rel), kb, tipo, kb * 1024, "✓ ya legible"))
                n_extraidos += 1
            elif tipo in ("IMAGEN", "VIDEO", "AUDIO", "EXCEL", "CSV"):
                inventario.append((sub, str(rel), kb, tipo, 0, "— se lista, no se extrae"))
            else:
                inventario.append((sub, str(rel), kb, tipo, 0, "— no extraíble"))

            print(f"[{sub}] {rel} ({kb} KB) → {inventario[-1][5]}")

    # Escribir inventario.md
    md = ["# Inventario del encargo\n"]
    md.append(f"Total archivos detectados: **{n_total}** ({n_extraidos} con texto disponible)\n")
    md.append("Generado por `extraer_textos.py`. Recorre recursivamente todas las subcarpetas\n")
    md.append("salvo `04_inspeccion/`, `99_salida/` y `_skill_workspace/`.\n")
    md.append("")
    md.append("| Carpeta raíz | Ruta relativa | KB | Tipo | Chars texto | Estado |")
    md.append("|--------------|---------------|---:|------|-----------:|--------|")
    for sub, rel, kb, tipo, chars, tag in inventario:
        md.append(f"| {sub} | `{rel}` | {kb} | {tipo} | {chars} | {tag} |")
    md.append("")
    md.append("## Cómo usar este inventario")
    md.append("")
    md.append("- Los archivos `.txt` paralelos a cada PDF/DOCX contienen el texto extraído. Léelos con Read tool.")
    md.append("- Los PDFs escaneados sin texto (estado ⚠) tienen un `.preview.png` para una mirada rápida; pide al usuario que los re-procese con OCR si son importantes.")
    md.append("- Las imágenes/vídeos/audios sólo se listan para que sepas que existen; abrirlos requiere herramientas específicas.")
    md.append("")
    (workspace / "inventario.md").write_text("\n".join(md), encoding="utf-8")
    print()
    print(f"Inventario escrito en: {workspace / 'inventario.md'}")
    print(f"Total: {n_total} archivos, {n_extraidos} con texto extraíble.")
    return 0


if __name__ == "__main__":
    sys.exit(main(sys.argv[1:]))
