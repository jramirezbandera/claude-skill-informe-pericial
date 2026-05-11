"""Ensambla el .docx v1 del informe pericial.

NO redacta prosa: sólo monta el documento desde piezas que la skill (Claude
en conversación) ya ha generado en `_skill_workspace/`.

Lee:
- <encargo>/caso.yaml
- <encargo>/deficiencias.md (la lista cerrada por /informe deficiencias)
- <encargo>/00_encargo/objeto_pericial.md
- <encargo>/04_inspeccion/fecha_inspeccion.txt
- <encargo>/04_inspeccion/fotos_renombradas/  (P001_DEFNN_slug.jpg)
- <encargo>/_skill_workspace/redaccion_deficiencias.md  (prosa por deficiencia, generada por la skill)
- <encargo>/_skill_workspace/documentacion_consultada.md  (prosa, generada por la skill)
- <encargo>/_skill_workspace/antecedentes.md  (opcional, prosa, generada por la skill)
- <skill>/boilerplate/{juramento_objetividad.md, metodologia.md, pie_firma.md, datos_autor.yaml}
- <skill>/plantillas/plantilla_base.docx

Escribe:
- <encargo>/99_salida/informe_pericial_v1.docx

Uso:
    python redactar_v1.py --encargo "<ruta-encargo>"

    # opcionalmente fija el lugar y la fecha de entrega para el pie de firma:
    python redactar_v1.py --encargo "<...>" --lugar "Málaga" --fecha-entrega 2026-05-20
"""
from __future__ import annotations

import argparse
import datetime as dt
import re
import sys
from pathlib import Path

import yaml
from docx import Document
from docx.shared import Cm, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from jinja2 import Template

try:
    sys.stdout.reconfigure(encoding="utf-8")
    sys.stderr.reconfigure(encoding="utf-8")
except Exception:
    pass

SKILL_ROOT = Path(__file__).resolve().parent.parent

MESES_ES = ["", "enero", "febrero", "marzo", "abril", "mayo", "junio",
            "julio", "agosto", "septiembre", "octubre", "noviembre", "diciembre"]

EXT_FOTO = {".jpg", ".jpeg", ".png", ".heic", ".heif", ".webp"}


def _fmt_fecha_larga(iso: str) -> str:
    """'2026-05-20' -> '20 de mayo de 2026'."""
    try:
        d = dt.date.fromisoformat(iso)
        return f"{d.day} de {MESES_ES[d.month]} de {d.year}"
    except Exception:
        return iso


def _render(template_str: str, ctx: dict) -> str:
    return Template(template_str, keep_trailing_newline=False).render(**ctx)


# ────────── Helpers de docx ──────────

def _add_heading(doc, texto: str, level: int) -> None:
    h = doc.add_paragraph(texto)
    h.style = doc.styles[f"Heading {level}"]


def _add_para(doc, texto: str, style: str = "Body Text") -> None:
    """Añade un párrafo. Acepta '\\n\\n' para dividir en varios."""
    if not texto:
        return
    for bloque in re.split(r"\n\s*\n", texto.strip()):
        bloque = bloque.strip()
        if not bloque:
            continue
        p = doc.add_paragraph()
        try:
            p.style = doc.styles[style]
        except KeyError:
            pass
        # Soporte mínimo de **negrita**
        for parte in re.split(r"(\*\*[^*]+\*\*)", bloque):
            if parte.startswith("**") and parte.endswith("**"):
                run = p.add_run(parte[2:-2])
                run.bold = True
            else:
                p.add_run(parte)


def _add_picture_with_caption(doc, path: Path, num: int, descripcion: str, width_cm: float = 14.0) -> None:
    """Inserta foto centrada + caption 'Imagen NN: descripción'."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(str(path), width=Cm(width_cm))
    cap = doc.add_paragraph()
    try:
        cap.style = doc.styles["Caption"]
    except KeyError:
        pass
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.add_run(f"Imagen {num}: {descripcion}")


def _add_page_break(doc) -> None:
    p = doc.add_paragraph()
    r = p.add_run()
    br = OxmlElement("w:br")
    br.set(qn("w:type"), "page")
    r._r.append(br)


def _add_toc_placeholder(doc) -> None:
    """Inserta un campo TOC. Word lo poblará al abrir el doc."""
    p = doc.add_paragraph()
    r = p.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = r' TOC \o "1-3" \h \z \u '
    fld_sep = OxmlElement("w:fldChar")
    fld_sep.set(qn("w:fldCharType"), "separate")
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    placeholder = OxmlElement("w:r")
    pt = OxmlElement("w:t")
    pt.text = "Pulsa F9 en Word para actualizar este índice."
    placeholder.append(pt)
    r._r.append(fld_begin)
    r._r.append(instr)
    r._r.append(fld_sep)
    r._r.append(placeholder)
    r._r.append(fld_end)


# ────────── Conversión HEIC ──────────

def _ensure_jpg(src: Path, workspace: Path) -> Path:
    """Si HEIC/HEIF, convierte a JPG en workspace/jpegs/. Devuelve ruta usable."""
    if src.suffix.lower() not in {".heic", ".heif"}:
        return src
    try:
        from PIL import Image
        import pillow_heif
        pillow_heif.register_heif_opener()
    except Exception:
        print(f"WARN: no se puede convertir HEIC ({src.name}), pillow-heif no disponible", file=sys.stderr)
        return src
    out_dir = workspace / "jpegs"
    out_dir.mkdir(parents=True, exist_ok=True)
    out = out_dir / (src.stem + ".jpg")
    if out.exists() and out.stat().st_mtime >= src.stat().st_mtime:
        return out
    with Image.open(src) as img:
        rgb = img.convert("RGB")
        rgb.save(out, "JPEG", quality=85)
    return out


# ────────── Parseo de los .md de la skill ──────────

def _parse_redaccion_deficiencias(md: str) -> dict[str, dict[str, str]]:
    """De redaccion_deficiencias.md saca {DEFNN: {existencia, definicion, causa, certificacion, propuesta}}.

    Formato esperado:
        ## DEF01 — slug · Título legible
        ### Existencia
        prosa…
        ### Definición en proyecto
        prosa…
        ### Causa
        prosa…
        ### Certificación   (opcional)
        prosa…
        ### Propuesta de reparación
        prosa…
    """
    SUBHEADS = {
        "existencia": "existencia",
        "definicion en proyecto": "definicion",
        "definición en proyecto": "definicion",
        "definicion en proyecto y mediciones": "definicion",
        "definición en proyecto y mediciones": "definicion",
        "causa": "causa",
        "certificacion": "certificacion",
        "certificación": "certificacion",
        "propuesta de reparacion": "propuesta",
        "propuesta de reparación": "propuesta",
    }
    out: dict[str, dict[str, str]] = {}
    current_def: str | None = None
    current_sub: str | None = None
    buf: list[str] = []

    def flush():
        if current_def and current_sub:
            out.setdefault(current_def, {})[current_sub] = "\n".join(buf).strip()

    for line in md.splitlines():
        m_def = re.match(r"^##\s+(DEF\d{1,3})\b", line.strip(), re.IGNORECASE)
        m_sub = re.match(r"^###\s+(.+?)\s*$", line.strip())
        if m_def:
            flush()
            buf = []
            current_def = m_def.group(1).upper()
            current_sub = None
            continue
        if m_sub and current_def:
            flush()
            buf = []
            sub_norm = m_sub.group(1).strip().lower()
            current_sub = SUBHEADS.get(sub_norm)
            continue
        if current_sub:
            buf.append(line)
    flush()
    return out


def _parse_deficiencias_md(md: str) -> list[dict]:
    """Parsea deficiencias.md (formato del subcomando /informe deficiencias)."""
    items = []
    current: dict | None = None
    for line in md.splitlines():
        m = re.match(r"^\s*(\d+)\.\s+\*\*([a-z0-9_]+)\*\*\s*·\s*(.+?)\s*$", line)
        if m:
            if current:
                items.append(current)
            current = {
                "id": int(m.group(1)),
                "slug": m.group(2),
                "titulo": m.group(3).strip(),
            }
    if current:
        items.append(current)
    return items


def _fotos_de_deficiencia(fotos_dir: Path, def_id: int) -> list[Path]:
    if not fotos_dir.exists():
        return []
    pat = re.compile(rf"^P\d+_DEF{def_id:02d}_", re.IGNORECASE)
    fotos = [p for p in sorted(fotos_dir.iterdir())
             if p.is_file() and p.suffix.lower() in EXT_FOTO and pat.match(p.name)]
    return fotos


def _descripcion_foto(p: Path, descripciones_json: dict | None = None) -> str:
    """Prioridad: descripción de fotos_descripciones.json (texto libre del usuario).
    Fallback: slug del nombre de archivo.
    """
    if descripciones_json:
        # Buscar por P_id (la clave canónica) y por archivo_final
        m_id = re.match(r"^(P\d+)", p.name, re.IGNORECASE)
        if m_id:
            p_id = m_id.group(1).upper()
            entry = descripciones_json.get(p_id)
            if entry and entry.get("descripcion"):
                return entry["descripcion"]
    m = re.match(r"^P\d+_(?:DEF\d+|GEN)_(.+?)\.[^.]+$", p.name, re.IGNORECASE)
    if not m:
        return p.stem
    return m.group(1).replace("_", " ").replace("-", " ")


# ────────── Secciones del informe ──────────

def seccion_consideraciones(doc, caso: dict, objeto: str) -> None:
    _add_heading(doc, "1. CONSIDERACIONES PREVIAS Y OBJETO DE TRABAJO", 1)
    autor = caso["_autor"]
    cliente = caso.get("cliente", {})
    nombre_cliente = cliente.get("nombre", "")
    rol = cliente.get("rol", "")
    intro = (
        f"El que suscribe, {autor['nombre_completo']}, {autor['titulacion']} colegiado nº {autor['colegiado']} "
        f"del {autor['colegio']}, recibe encargo de "
        f"{nombre_cliente}{' (parte ' + rol + ')' if rol else ''} "
        f"para emitir el presente dictamen pericial."
    )
    _add_para(doc, intro)
    if objeto:
        _add_para(doc, "El objeto del presente trabajo es:")
        _add_para(doc, objeto)


def seccion_juramento(doc) -> None:
    _add_heading(doc, "2. JURAMENTO DE ACTUAR CON OBJETIVIDAD", 1)
    juramento = (SKILL_ROOT / "boilerplate" / "juramento_objetividad.md").read_text(encoding="utf-8")
    _add_para(doc, juramento)


def seccion_bases(doc, caso: dict, doc_consultada: str) -> None:
    _add_heading(doc, "3. BASES DEL DICTAMEN. VISITA E INSPECCIÓN", 1)

    _add_heading(doc, "3.1. Visita de inspección", 2)
    fecha_visita = caso.get("fechas", {}).get("visita") or caso.get("_fecha_visita") or ""
    if fecha_visita:
        _add_para(doc, f"La visita de inspección al inmueble se realizó el día {_fmt_fecha_larga(fecha_visita)}, "
                       f"durante la cual se llevó a cabo el reconocimiento visual de los elementos objeto del "
                       f"presente dictamen y se obtuvo el reportaje fotográfico que se incorpora al informe.")
    else:
        _add_para(doc, "[[FALTA: fecha de inspección]]")

    _add_heading(doc, "3.2. Documentación examinada", 2)
    if doc_consultada:
        _add_para(doc, doc_consultada)
    else:
        _add_para(doc, "[[FALTA: documentación consultada — generar en _skill_workspace/documentacion_consultada.md]]")

    _add_heading(doc, "3.3. Alcance de esta base documental", 2)
    _add_para(doc,
        "El presente dictamen se emite a partir de la documentación relacionada en el apartado precedente "
        "y de las observaciones realizadas durante la visita de inspección. Las conclusiones que se "
        "exponen lo son sin perjuicio de aquellos extremos que pudieran resultar de información o "
        "documentación adicional que no haya sido facilitada al suscrito."
    )


def seccion_antecedentes(doc, caso: dict, antecedentes_md: str) -> None:
    _add_heading(doc, "4. ANTECEDENTES Y DESCRIPCIÓN DEL CONJUNTO", 1)
    if antecedentes_md.strip():
        _add_para(doc, antecedentes_md)
        return

    # Fallback estructurado a partir de caso.yaml
    inm = caso.get("inmueble", {})
    ag = caso.get("agentes_obra", {})

    _add_heading(doc, "4.1. Situación", 2)
    direccion = inm.get("direccion", "")
    _add_para(doc, f"El inmueble objeto del dictamen se encuentra en {direccion or '[[FALTA: dirección]]'}.")
    if inm.get("ref_catastral"):
        _add_para(doc, f"Referencia catastral: {inm['ref_catastral']}.")

    _add_heading(doc, "4.2. Fechas de obra y agentes que intervinieron", 2)
    lineas = []
    for clave, etiqueta in [("promotor", "Promotor"), ("constructor", "Constructor"),
                             ("proyectista", "Proyectista"), ("direccion_facultativa", "Dirección facultativa"),
                             ("direccion_ejecucion", "Dirección de ejecución de obra")]:
        if ag.get(clave):
            lineas.append(f"**{etiqueta}**: {ag[clave]}.")
    if ag.get("fecha_inicio_obra"):
        lineas.append(f"**Inicio de obra**: {_fmt_fecha_larga(ag['fecha_inicio_obra'])}.")
    if ag.get("fecha_fin_obra"):
        lineas.append(f"**Fin de obra**: {_fmt_fecha_larga(ag['fecha_fin_obra'])}.")
    if ag.get("fecha_recepcion"):
        lineas.append(f"**Recepción**: {_fmt_fecha_larga(ag['fecha_recepcion'])}.")
    if not lineas:
        lineas.append("[[FALTA: agentes y fechas de obra]]")
    for l in lineas:
        _add_para(doc, l)

    _add_heading(doc, "4.3. Ámbito de la prueba pericial solicitada", 2)
    _add_para(doc, caso.get("objeto_pericial") or "[[FALTA: objeto pericial]]")

    _add_heading(doc, "4.4. Tipología y entorno", 2)
    if inm.get("tipologia"):
        _add_para(doc, f"Tipología: {inm['tipologia']}.")
    if inm.get("superficie_construida_m2"):
        _add_para(doc, f"Superficie construida: {inm['superficie_construida_m2']} m².")


def seccion_metodologia(doc, caso: dict) -> None:
    if not caso.get("secciones_a_incluir", {}).get("metodologia", True):
        return
    _add_heading(doc, "5. METODOLOGÍA DE ANÁLISIS", 1)
    tpl = (SKILL_ROOT / "boilerplate" / "metodologia.md").read_text(encoding="utf-8")
    ctx = {
        "hay_informe_contraria": caso.get("secciones_a_incluir", {}).get("analisis_critico_contraria", False),
        "proyectista": caso.get("agentes_obra", {}).get("proyectista", ""),
    }
    _add_para(doc, _render(tpl, ctx))


def seccion_deficiencias(doc, caso: dict, deficiencias: list[dict],
                          redaccion: dict, fotos_dir: Path, workspace: Path,
                          fotos_descripciones: dict | None = None) -> None:
    _add_heading(doc, "7. DEFICIENCIAS OBSERVADAS", 1)
    contador_imagenes = 0

    for d in deficiencias:
        defcod = f"DEF{d['id']:02d}"
        titulo = f"7.{d['id']}. {d['titulo'].upper()}"
        _add_heading(doc, titulo, 2)

        prosa = redaccion.get(defcod, {})

        # 7.x.1 Existencia + fotos
        _add_heading(doc, f"7.{d['id']}.1. Existencia", 3)
        if prosa.get("existencia"):
            _add_para(doc, prosa["existencia"])
        else:
            _add_para(doc, f"[[FALTA: redactar Existencia de {defcod} en _skill_workspace/redaccion_deficiencias.md]]")

        fotos = _fotos_de_deficiencia(fotos_dir, d["id"])
        for foto in fotos:
            jpg = _ensure_jpg(foto, workspace)
            contador_imagenes += 1
            _add_picture_with_caption(doc, jpg, contador_imagenes, _descripcion_foto(foto, fotos_descripciones))

        # 7.x.2 Definición en proyecto y mediciones
        _add_heading(doc, f"7.{d['id']}.2. Definición en proyecto y mediciones", 3)
        if prosa.get("definicion"):
            _add_para(doc, prosa["definicion"])
        else:
            _add_para(doc, f"[[FALTA: redactar Definición en proyecto de {defcod}]]")

        # 7.x.3 Causa
        _add_heading(doc, f"7.{d['id']}.3. Causa", 3)
        if prosa.get("causa"):
            _add_para(doc, prosa["causa"])
        else:
            _add_para(doc, f"[[FALTA: redactar Causa de {defcod}]]")

        # 7.x.4 Certificación (opcional)
        if prosa.get("certificacion"):
            _add_heading(doc, f"7.{d['id']}.4. Certificación", 3)
            _add_para(doc, prosa["certificacion"])
            offset = 1
        else:
            offset = 0

        # 7.x.5 Propuesta de reparación
        _add_heading(doc, f"7.{d['id']}.{4 + offset}. Propuesta de reparación", 3)
        if prosa.get("propuesta"):
            _add_para(doc, prosa["propuesta"])
        else:
            _add_para(doc, f"[[FALTA: redactar Propuesta de reparación de {defcod}]]")
        _add_para(doc, f"[[PRESUPUESTO_DEFICIENCIA: {d['slug']}]]")


def seccion_cuadro_resumen(doc, deficiencias: list[dict]) -> None:
    _add_heading(doc, "10. CUADRO RESUMEN Y CONCLUSIONES", 1)
    _add_para(doc, "Se relacionan a continuación las deficiencias analizadas, junto con su causa "
                   "atribuida y el coste estimado de reparación:")

    table = doc.add_table(rows=1, cols=5)
    table.style = "Table Grid" if "Table Grid" in [s.name for s in doc.styles] else None
    hdr = table.rows[0].cells
    for i, h in enumerate(["Nº", "Deficiencia", "Causa breve", "Imputable a", "Coste reparación"]):
        hdr[i].text = h
    for d in deficiencias:
        row = table.add_row().cells
        row[0].text = str(d["id"])
        row[1].text = d["titulo"]
        row[2].text = "[[CAUSA BREVE]]"
        row[3].text = "[[IMPUTABLE A]]"
        row[4].text = "[[COSTE]]"


def seccion_pie_firma(doc, caso: dict, lugar: str, fecha_entrega: str) -> None:
    _add_heading(doc, "12. FIRMA", 1)
    tpl = (SKILL_ROOT / "boilerplate" / "pie_firma.md").read_text(encoding="utf-8")
    autor = caso.get("_autor", {})
    ctx = {
        "lugar": lugar or autor.get("despacho", {}).get("ciudad", ""),
        "fecha_entrega_largo": _fmt_fecha_larga(fecha_entrega) if fecha_entrega else _fmt_fecha_larga(dt.date.today().isoformat()),
        "nombre_completo": autor.get("nombre_completo", ""),
        "titulacion": autor.get("titulacion", ""),
        "colegiado": autor.get("colegiado", ""),
        "colegio": autor.get("colegio", ""),
        "despacho": autor.get("despacho", {}),
        "email": autor.get("email", ""),
    }
    _add_para(doc, _render(tpl, ctx))


# ────────── Cabecera del header (sustituir titulo_corto) ──────────

def _patch_header(doc, titulo_corto: str) -> None:
    """Sustituye el marcador {{titulo_corto}} en el header por el valor real."""
    if not titulo_corto:
        return
    for sec in doc.sections:
        for para in sec.header.paragraphs:
            for run in para.runs:
                if "{{titulo_corto}}" in run.text:
                    run.text = run.text.replace("{{titulo_corto}}", titulo_corto)


# ────────── Main ──────────

def main(argv: list[str]) -> int:
    p = argparse.ArgumentParser()
    p.add_argument("--encargo", required=True)
    p.add_argument("--lugar", default="Málaga")
    p.add_argument("--fecha-entrega", default=None, help="ISO YYYY-MM-DD; default hoy")
    args = p.parse_args(argv)

    encargo = Path(args.encargo).resolve()
    if not encargo.exists():
        print(f"ERROR: no existe {encargo}", file=sys.stderr)
        return 1

    # Cargar caso.yaml
    caso_path = encargo / "caso.yaml"
    if not caso_path.exists():
        print(f"ERROR: no existe {caso_path}", file=sys.stderr)
        return 1
    caso = yaml.safe_load(caso_path.read_text(encoding="utf-8")) or {}

    # Cargar autor
    autor = yaml.safe_load((SKILL_ROOT / "boilerplate" / "datos_autor.yaml").read_text(encoding="utf-8"))
    caso["_autor"] = autor

    # Objeto pericial
    obj_path = encargo / "00_encargo" / "objeto_pericial.md"
    objeto = obj_path.read_text(encoding="utf-8").strip() if obj_path.exists() else caso.get("objeto_pericial", "")

    # Fecha de inspección
    fecha_path = encargo / "04_inspeccion" / "fecha_inspeccion.txt"
    if fecha_path.exists() and not caso.get("fechas", {}).get("visita"):
        fecha_txt = fecha_path.read_text(encoding="utf-8").strip().split(" a ")[0]
        caso.setdefault("fechas", {})["visita"] = fecha_txt
        caso["_fecha_visita"] = fecha_txt

    # Workspace
    ws = encargo / "_skill_workspace"
    ws.mkdir(parents=True, exist_ok=True)

    # Documentación consultada (la skill la genera en sesión)
    doc_consultada_path = ws / "documentacion_consultada.md"
    doc_consultada = doc_consultada_path.read_text(encoding="utf-8").strip() if doc_consultada_path.exists() else ""

    # Antecedentes (opcional, la skill los genera si quiere prosa propia)
    antecedentes_path = ws / "antecedentes.md"
    antecedentes = antecedentes_path.read_text(encoding="utf-8").strip() if antecedentes_path.exists() else ""

    # Redacción de deficiencias
    redaccion_path = ws / "redaccion_deficiencias.md"
    redaccion = {}
    if redaccion_path.exists():
        redaccion = _parse_redaccion_deficiencias(redaccion_path.read_text(encoding="utf-8"))

    # Lista de deficiencias
    def_path = encargo / "deficiencias.md"
    deficiencias = []
    if def_path.exists():
        deficiencias = _parse_deficiencias_md(def_path.read_text(encoding="utf-8"))
    if not deficiencias:
        print("ERROR: deficiencias.md vacío o ausente. Ejecuta /informe deficiencias antes.", file=sys.stderr)
        return 1

    # Fotos
    fotos_dir = encargo / "04_inspeccion" / "fotos_renombradas"
    fotos_descripciones = None
    fd_path = ws / "fotos_descripciones.json"
    if fd_path.exists():
        try:
            import json as _json
            fotos_descripciones = _json.loads(fd_path.read_text(encoding="utf-8"))
        except Exception as e:
            print(f"WARN: no se pudo leer fotos_descripciones.json: {e}", file=sys.stderr)

    # Abrir plantilla
    plantilla = SKILL_ROOT / "plantillas" / "plantilla_base.docx"
    doc = Document(str(plantilla))

    # Header con titulo_corto
    titulo_corto = caso.get("titulo_corto") or ""
    if not titulo_corto and caso.get("cliente", {}).get("nombre"):
        titulo_corto = f"{caso['cliente']['nombre']} — {caso.get('inmueble', {}).get('direccion', '').split(',')[-1].strip() or ''}".strip(" —")
    _patch_header(doc, titulo_corto)

    # Portada
    portada = doc.add_paragraph()
    portada.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = portada.add_run("DICTAMEN PERICIAL")
    r.bold = True
    r.font.size = Pt(22)
    if titulo_corto:
        sub = doc.add_paragraph()
        sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = sub.add_run(titulo_corto)
        r.font.size = Pt(14)
    autor_p = doc.add_paragraph()
    autor_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    autor_p.add_run(f"\n{autor['nombre_completo']}\n{autor['titulacion']} — Colegiado nº {autor['colegiado']}\n{autor['colegio']}")
    fecha_p = doc.add_paragraph()
    fecha_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    f_entrega = args.fecha_entrega or dt.date.today().isoformat()
    fecha_p.add_run(_fmt_fecha_larga(f_entrega))
    _add_page_break(doc)

    # Índice
    _add_heading(doc, "ÍNDICE", 1)
    _add_toc_placeholder(doc)
    _add_page_break(doc)

    # Cuerpo
    seccion_consideraciones(doc, caso, objeto)
    seccion_juramento(doc)
    seccion_bases(doc, caso, doc_consultada)
    seccion_antecedentes(doc, caso, antecedentes)
    seccion_metodologia(doc, caso)

    # 6. Desarrollo anormal de la obra (opcional)
    if caso.get("secciones_a_incluir", {}).get("desarrollo_anormal_obra"):
        _add_heading(doc, "6. DESARROLLO ANORMAL DE LA OBRA", 1)
        anormal_md = (ws / "desarrollo_anormal.md")
        if anormal_md.exists():
            _add_para(doc, anormal_md.read_text(encoding="utf-8"))
        else:
            _add_para(doc, "[[FALTA: redactar desarrollo_anormal.md]]")

    seccion_deficiencias(doc, caso, deficiencias, redaccion, fotos_dir, ws, fotos_descripciones)

    # 8. Análisis crítico contraria (opcional)
    if caso.get("secciones_a_incluir", {}).get("analisis_critico_contraria"):
        _add_heading(doc, "8. ANÁLISIS CRÍTICO DEL INFORME DE LA ACTORA", 1)
        critico_md = (ws / "analisis_critico_contraria.md")
        if critico_md.exists():
            _add_para(doc, critico_md.read_text(encoding="utf-8"))
        else:
            _add_para(doc, "[[FALTA: redactar analisis_critico_contraria.md]]")

    # 9. Penalizaciones (opcional)
    if caso.get("secciones_a_incluir", {}).get("penalizaciones_retraso"):
        _add_heading(doc, "9. PENALIZACIONES Y RETRASO", 1)
        pen_md = (ws / "penalizaciones.md")
        if pen_md.exists():
            _add_para(doc, pen_md.read_text(encoding="utf-8"))
        else:
            _add_para(doc, "[[FALTA: redactar penalizaciones.md]]")

    seccion_cuadro_resumen(doc, deficiencias)

    # 11. Hoja resumen presupuesto (opcional)
    if caso.get("secciones_a_incluir", {}).get("hoja_resumen_presupuesto", True):
        _add_heading(doc, "11. HOJA RESUMEN DE PRESUPUESTO", 1)
        _add_para(doc, "[[HOJA_RESUMEN_PRESUPUESTO — se inyecta en /informe presupuestar]]")

    seccion_pie_firma(doc, caso, args.lugar, f_entrega)

    # Guardar
    out_dir = encargo / "99_salida"
    out_dir.mkdir(parents=True, exist_ok=True)
    out = out_dir / "informe_pericial_v1.docx"
    doc.save(str(out))

    print(f"OK — informe v1 generado: {out}")
    print(f"     Tamaño: {out.stat().st_size/1024:.0f} KB")
    print(f"     Deficiencias procesadas: {len(deficiencias)}")
    n_falta = sum(1 for d in deficiencias if not redaccion.get(f"DEF{d['id']:02d}"))
    if n_falta:
        print(f"     AVISO: {n_falta}/{len(deficiencias)} deficiencias sin prosa redactada (marcadas [[FALTA: ...]])")
    if not doc_consultada:
        print(f"     AVISO: documentacion_consultada.md no existe (sección 3.2 con marcador)")
    return 0


if __name__ == "__main__":
    sys.exit(main(sys.argv[1:]))
